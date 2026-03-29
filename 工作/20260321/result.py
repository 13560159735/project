import pandas as pd
from datetime import datetime
import json
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
import xlwings as xw
import shutil
import requests
import os
import openpyxl
from docx import Document
from collections import deque
from docx.shared import Pt
from docx.oxml.ns import qn

word_path="1.docx"
doc=Document(word_path)
tables=doc.tables

CACHE_FILE='cache.json'

def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

cache=load_cache()

def fetch_data(url):
    print('url',url)
    if url in cache:
        print('走缓存')
        return cache[url]
    print('调用接口')
    headers={'Authorization': '28041b0a-feec-49df-9d42-21c21e208bca'}
    res=requests.get(url,headers=headers)
    response=res.json
    cache[url]=response
    save_cache(cache)
    return response

res=fetch_data('http://open.api.tianyancha.com/services/v3/open/investtree?flag=4&dir=down&keyword=深圳市前海一方科技研发集团有限公司&minPercent=0&maxPercent=1')
investtree=json.loads(res["result"])

rankToChinese={
    1:'一级',
    2:'二级',
    3:'三级',
    4:'四级',
    5:'五级',
    6:'六级',
    7:'七级',
    8:'八级',
    9:'九级',
    10:'十级',
}

queue=deque([(investtree[0],0)])
investtreeList=[]
while queue:
    node,level=queue.popleft()
    if level !=0 and node['regStatus']=='存续':
        currentType='控股公司'
        if level == 1:
            currentType='控股子公司'
        percent=f"{float(node['percent'])*100:.2f}"
        investtreeList.append({
            '子公司名称':node['name'],
            '子公司类型':currentType,
            '级次':rankToChinese[level],
            '持股比例（%）':percent,
            '表决权比例（%）':percent,
            '状态':node['regStatus'],
            '社会统一信用代码':node['creditCode'],

        })
    for index,item in enumerate(node['children']):
        queue.append((item,level+1))
with open('investtreeList1.json', 'w', encoding='utf-8') as f:
    json.dump(investtreeList, f, ensure_ascii=False, indent=4)

for index,item in enumerate(investtreeList):
    if float(item['持股比例（%）'])<40:
        currentRes=fetch_data(f"http://open.api.tianyancha.com/services/open/ic/baseinfoV3/2.0?keyword={item['社会统一信用代码']}")
        item['注册地']=currentRes['result']['city']
        item['业务性质']=currentRes['result']['industryAll']['category']
        item['对合营企业或联营企业投资的会计处理方法']='权益法'
        item['注册资本（万元）']=f"{float(currentRes['result']['regCapital'].split('万')[0]):,.2f}"
with open('investtreeList2.json', 'w', encoding='utf-8') as f:
    json.dump(investtreeList, f, ensure_ascii=False, indent=4)
        

for i in range(len(tables[1].rows)-1,0,-1):
    tbl=tables[1]._tbl
    tr=tables[1].rows[i]._tr
    tbl.remove(tr)

for i in range(len(tables[2].rows)-1,1,-1):
    tbl=tables[2]._tbl
    tr=tables[2].rows[i]._tr
    tbl.remove(tr)

for i in range(len(tables[3].rows)-1,0,-1):
    tbl=tables[3]._tbl
    tr=tables[3].rows[i]._tr
    tbl.remove(tr)



for index,item in enumerate(investtreeList):
    if float(item['持股比例（%）'])<40:
        row=tables[2].add_row().cells
        valueList=[
            item['子公司名称'],
            item['注册地'],
            item['业务性质'],
            item['持股比例（%）'],
            '',
            item['表决权比例（%）'],
            item['对合营企业或联营企业投资的会计处理方法'],
            item['注册资本（万元）'],
        ]
        for indexValueList,itemValueList in enumerate(valueList):
            p=row[indexValueList].paragraphs[0]
            run=p.add_run(itemValueList)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    else:
        row=tables[1].add_row().cells
        valueList=[
            item['子公司名称'],
            item['子公司类型'],
            item['级次'],
            item['持股比例（%）'],
            item['表决权比例（%）'],
        ]

        for indexValueList,itemValueList in enumerate(valueList):
            p=row[indexValueList].paragraphs[0]
            run=p.add_run(itemValueList)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

doc.save("1_填充后.docx")
