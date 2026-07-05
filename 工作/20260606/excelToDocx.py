from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook

def is_blank_row(row_values):
    return all(value is None or str(value).strip()=='' for value in row_values)

def format_value_for_doc(value):
    if value is None:
        return ""
    if isinstance(value,bool):
        return str(value)
    if isinstance(value,(int,float)):
        return f'{value:,.2f}'
    text=str(value).strip()
    if text == '':
        return ''
    try:
        num=float(text.replace(',',''))
        return f'{num:,.2f}'
    except ValueError:
        return text

def set_paragraph_font(para):
    for run in para.runs:
        run.font.name='宋体'
        run.font.size=Pt(10.5)
        rfonts = run._element.rPr.rFonts
        rfonts.set(qn("w:eastAsia"), "宋体")

def set_cell_font(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name='宋体'
            run.font.size=Pt(10.5)
            rfonts = run._element.rPr.rFonts
            rfonts.set(qn("w:eastAsia"), "宋体")

def iter_doc_elements_from_title(doc,target_title):
    started=False
    for child in doc.element.body:
        if child.tag.endswith('p'):
            para=next((p for p in doc.paragraphs if p._element is child),None)
            if para is None:
                continue
            if not started and target_title in para.text:
                started = True
                yield 'p',para
            elif started:
                yield 'p',para
        
        elif started and child.tag.endswith('tbl'):
            table=next((t for t in doc.tables if t._element is child),None)
            if table is not None:
                yield 'tbl',table







                    











