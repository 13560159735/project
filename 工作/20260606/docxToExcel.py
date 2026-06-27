import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, numbers, Font


def try_convert_to_number(s):
    """将可能含千分位逗号的字符串转为数字，成功返回(数字,True)，否则返回(None,False)"""
    if not isinstance(s, str):
        return s, False
    s = s.strip()
    if s == "":
        return None, False
    s_clean = s.replace(',', '')
    if re.match(r'^-?\d+(?:\.\d+)?$', s_clean):
        try:
            num = float(s_clean)
            if num.is_integer():
                return int(num), True
            else:
                return num, True
        except ValueError:
            return None, False
    return None, False

def paragraph_has_bold(para):
    """判断段落中是否有加粗的文本"""
    for run in para.runs:
        if run.bold:
            return True
    return False


def get_cell_merge_info(table):
    merge_info=[]
    processed_cells=set()

    rows=table.rows


    for r , row in enumerate(rows):
        for c , cell in enumerate(row.cells):
            if (r,c) in processed_cells:
                continue
            tcPr=cell._element.get_or_add_tcPr()
            grid_span = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            h_span = int(grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if grid_span is not None else 1
            
            # 检查垂直合并（vMerge）
            v_merge = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
            v_span=1
            if v_merge is not None:
                for check_r in range(r+1,len(rows)):
                    # check_cell=rows[check_r].cells[c] if c<len(rows[check_r].cells) else None
                    check_cell= None
                    if c<len(rows[check_r].cells):
                        check_cell=rows[check_r].cells[c]

                    if check_cell is None:
                        break
                    check_tcPr=check_cell._element.get_or_add_tcPr()
                    check_v_merge = check_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                    if check_v_merge is not None:
                        v_span+=1
                    else:
                        break
            if h_span > 1 or v_span > 1:
                merge_info.append((r,c,v_span,h_span))
                for mr in range(r,r+v_span):
                    for mc in range(c,c+h_span):
                        processed_cells.add((mr,mc))
    return merge_info

def extract_with_formatting(docx_path, target_title, output_excel):
    doc = Document(docx_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    found = False
    row_idx = 1

    table_start_rows=[]

    for element in doc.element.body:
        if element.tag.endswith('p'):          # 段落
            para = next((p for p in doc.paragraphs if p._element is element), None)
            if para is None:
                continue
            text = para.text.strip()
            if not found:
                if target_title in text:
                    found = True
                else:
                    continue
            cell = ws.cell(row=row_idx, column=1, value=text)
            # 如果段落中有加粗文本，则设置单元格字体加粗
            if paragraph_has_bold(para):
                cell.font = Font(bold=True)
            row_idx += 1

        elif element.tag.endswith('tbl'):       # 表格
            if not found:
                continue
            table = next((t for t in doc.tables if t._element is element), None)
            if table is None:
                continue

            start_row = row_idx
            merge_info=get_cell_merge_info(table)
            for r,row in enumerate(table.rows):
                for c,cell in enumerate(row.cells):
                    raw_text=cell.text.strip()
                    num_val,is_num=try_convert_to_number(raw_text)
                    excel_cell=ws.cell(row=start_row+r,column=1+c)
                    if is_num:
                        excel_cell.value = num_val
                        if isinstance(num_val, int):
                            excel_cell.number_format = '#,##0'
                        else:
                            excel_cell.number_format = '#,##0.00'
                    else:
                        excel_cell.value = raw_text
                    
                    has_bold=False
                    for para in cell.paragraphs:
                        if paragraph_has_bold(para):
                            has_bold=True
                            break
                    if has_bold:
                        excel_cell.font=Font(bold=True)

            end_row=start_row+len(table.rows)-1   
            max_col=max((len(row.cells) for row in table.rows) , default=0)
            table_start_rows.append((start_row,end_row,max_col,merge_info))   
            #创建一个字典，快速查找合并信息
            merge_dict={(r,c):(v_span,h_span) for r,c,v_span,h_span in merge_info}  
            top_bold=Side(style='medium')
            bottom_bold=Side(style='medium')
            thin_side=Side(style='thin')
            no_side=Side(style=None)
            for i in range(start_row,end_row+1):
                for j in range(1,max_col+1):
                    #获取相对表格的行列位置
                    rel_row=i-start_row
                    rel_col=j-1
                    #检查这个单元格是否为合并区域的一部分（但不是起始单元格）
                    is_merged_none_start=False
                    for r,c,v_span,h_span in merge_info:
                        if r <= rel_row < r+v_span and c <= rel_col < c+h_span and not(rel_row==r and rel_col==c):
                            is_merged_none_start=True
                            break
                    if is_merged_none_start:
                        continue
                    if (rel_row,rel_col) in merge_dict:
                        v_span,h_span = merge_dict[(rel_row,rel_col)]
                        #对于合并单元格，根据合并范围设置边框
                        top=top_bold if i == start_row else thin_side
                        bottom=bottom_bold if i+v_span-1 == end_row else thin_side
                        left=no_side if j == 1 else thin_side
                        right=no_side if j+h_span-1 == max_col else thin_side










