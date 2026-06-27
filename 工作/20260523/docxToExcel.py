import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, numbers, Font

def try_convert_to_number(s):
    """将可能含千分位逗号的字符串转为数字，成功返回(数字,True)，否则返回(None,False)"""
    if not isinstance(s, str):
        return s, False
    s = s.strip()
    if s == "":
        return None, False
    s_clean = s.replace(',', '')
    if re.match(r'^-?\d+(?:\.\d+)?$', s_clean):
        try:
            num = float(s_clean)
            if num.is_integer():
                return int(num), True
            else:
                return num, True
        except ValueError:
            return None, False
    return None, False

def paragraph_has_bold(para):
    """判断段落中是否有加粗的文本"""
    for run in para.runs:
        if run.bold:
            return True
    return False

def extract_with_formatting(docx_path, target_title, output_excel):
    doc = Document(docx_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    found = False
    row_idx = 1

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for element in doc.element.body:
        if element.tag.endswith('p'):          # 段落
            para = next((p for p in doc.paragraphs if p._element is element), None)
            if para is None:
                continue
            text = para.text.strip()
            if not found:
                if target_title in text:
                    found = True
                else:
                    continue
            cell = ws.cell(row=row_idx, column=1, value=text)
            # 如果段落中有加粗文本，则设置单元格字体加粗
            if paragraph_has_bold(para):
                cell.font = Font(bold=True)
            row_idx += 1

        elif element.tag.endswith('tbl'):       # 表格
            if not found:
                continue
            table = next((t for t in doc.tables if t._element is element), None)
            if table is None:
                continue

            start_row = row_idx
            # 写入表格数据，同时处理数字格式
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    raw_text = cell.text.strip()
                    num_val, is_num = try_convert_to_number(raw_text)
                    excel_cell = ws.cell(row=start_row + r, column=1 + c)
                    if is_num:
                        excel_cell.value = num_val
                        if isinstance(num_val, int):
                            excel_cell.number_format = '#,##0'
                        else:
                            excel_cell.number_format = '#,##0.00'
                    else:
                        excel_cell.value = raw_text
                    # 可选的：如果表格内需要加粗（例如表头），可在此处判断单元格中的段落是否有加粗
                    # 为保持简洁，本例不处理表格内加粗，如有需要可另行扩展
            # 加边框
            end_row = start_row + len(table.rows) - 1
            max_col = max((len(row.cells) for row in table.rows), default=0)
            for i in range(start_row, end_row + 1):
                for j in range(1, max_col + 1):
                    ws.cell(row=i, column=j).border = thin_border
            # 表格后空一行
            row_idx += len(table.rows) + 1

    wb.save(output_excel)
    print(f"已保存至: {output_excel}")

if __name__ == "__main__":
    docx_file = r"D:\宣达-深圳佑荣2025年审计报告\5、深圳佑荣审计报告-2025年单体\3、佑荣科技2025年财审报告附注.docx"
    excel_file = r"D:\宣达-深圳佑荣2025年审计报告\5、深圳佑荣审计报告-2025年单体\附注_合并项目注释_完整格式.xlsx"
    os.makedirs(os.path.dirname(excel_file), exist_ok=True)
    extract_with_formatting(docx_file, "八、财务报表主要项目注释", excel_file)